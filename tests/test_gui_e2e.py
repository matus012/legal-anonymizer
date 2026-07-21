"""Offscreen GUI tests. QT_QPA_PLATFORM=offscreen is set BEFORE QApplication import."""
import os

os.environ.setdefault("QT_QPA_PLATFORM", "offscreen")

import docx as _docx
import pytest
from PySide6.QtWidgets import QApplication

from gui.app import MainWindow


def _mk_docx(tmp_path, text):
    p = tmp_path / "in.docx"
    d = _docx.Document()
    d.add_paragraph(text)
    d.save(str(p))
    return str(p)


@pytest.fixture(scope="session")
def qapp():
    app = QApplication.instance() or QApplication([])
    yield app


def test_window_builds_with_three_pages(qapp):
    w = MainWindow()
    assert w.pages.count() == 3


def test_full_flow_scan_review_export(qapp, tmp_path):
    src = _mk_docx(tmp_path, "Jan Novak, DIC 2023456789, rodne cislo 835112/0009.")
    w = MainWindow()
    w.add_files([str(src)])
    w.known_edit.setPlainText("Jan Novak")
    w.start_scan(blocking=True)          # synchronous path for tests
    assert w.pages.currentIndex() == 1   # review page
    # auto rows ticked, review rows unticked:
    states = w.current_row_states()      # dict group -> bool
    scan = w.scans[str(src)]
    for r in scan.rows:
        assert states[r.group] == (r.bucket == "auto")
    w.start_export(blocking=True)
    assert w.pages.currentIndex() == 2   # done page
    out = tmp_path / "in_anon.docx"
    assert out.exists() and (tmp_path / "in_anon_report.txt").exists()
    txt = "\n".join(p.text for p in _docx.Document(str(out)).paragraphs)
    assert "Novak" not in txt

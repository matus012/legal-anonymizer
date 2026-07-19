"""W4a (context.md §10): accept all Word tracked revisions BEFORE any redaction runs, so that
deleted text is physically removed (never even a redaction candidate) and inserted text is
retained as ordinary body text — which then still flows through the normal W1-W3 redaction.

Tracked-change structure is ASYMMETRIC (verified against corpus bytes, kupna_zmluva_000.docx):
  * <w:ins ...><w:r><w:t>TEXT</w:t></w:r></w:ins>       -> accepted INSERT: unwrap, runs survive
  * <w:del ...><w:r><w:delText>TEXT</w:delText></w:r></w:del> -> DELETE still in file: drop subtree

Every fixture here is built entirely in-memory (this module must NOT import corpus/): a minimal
python-docx Document with <w:ins>/<w:del> spliced onto a <w:p> as raw XML, plus — for the
note-part coverage — a comments part (element path) and a footnotes part (blob path), each
hand-attached exactly as the W3 test does.

RED/GREEN: run this file against the pre-W4a writer and the assertions FAIL because the deleted
marker (and its <w:del>/<w:delText>) survives untouched into the output (the message prints the
surviving bytes). After wiring _strip_tracked_changes into redact_docx_body they pass: deleted
markers gone, no w:ins/w:del/w:delText remain, inserted text retained as ordinary <w:t>, an
insertion that is auto-detectable PII redacted, and every part reopens clean.
"""
from __future__ import annotations

import zipfile

from docx import Document
from docx.opc.constants import RELATIONSHIP_TYPE as RT
from docx.opc.packuri import PackURI
from docx.opc.part import Part
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls, qn

from writer.docx_body import redact_docx_body

INS_MARK = "INSERTEDKEEP"   # inside <w:ins>, must SURVIVE as ordinary body text
DEL_MARK = "DELETEDLEAK"     # inside <w:del><w:delText>, must be PHYSICALLY GONE
ICO = "48098191"            # checksum-valid ICO -> [ICO]; used to prove insertions still redact

_CM_CT = "application/vnd.openxmlformats-officedocument.wordprocessingml.comments+xml"
_FN_CT = "application/vnd.openxmlformats-officedocument.wordprocessingml.footnotes+xml"

# Footnote separator entries: <w:p> with no <w:t>, must survive verbatim (no-op through strip).
_FN_SEPS = (
    '<w:footnote w:type="separator" w:id="-1"><w:p><w:r><w:separator/></w:r></w:p></w:footnote>'
    '<w:footnote w:type="continuationSeparator" w:id="0"><w:p><w:r>'
    "<w:continuationSeparator/></w:r></w:p></w:footnote>"
)


def _ins_xml(text: str, wid: str = "9101") -> str:
    return (
        f'<w:ins {nsdecls("w")} w:id="{wid}" w:author="Advokat" '
        f'w:date="2024-01-01T00:00:00Z"><w:r>'
        f'<w:t xml:space="preserve">{text}</w:t></w:r></w:ins>'
    )


def _del_xml(text: str, wid: str = "9102", ns: bool = True) -> str:
    decl = f" {nsdecls('w')}" if ns else ""
    return (
        f'<w:del{decl} w:id="{wid}" w:author="Advokat" '
        f'w:date="2024-01-01T00:00:00Z"><w:r>'
        f'<w:delText xml:space="preserve">{text}</w:delText></w:r></w:del>'
    )


def _attach(doc, name: str, ct: str, xml: str, rt) -> None:
    """Attach ``xml`` as a generic Part at /word/<name> and relate it from the main part. On
    reopen comments+xml maps to a registered CommentsPart (element path); footnotes+xml has no
    registered class and stays a generic blob Part (blob path)."""
    part = Part(PackURI(f"/word/{name}"), ct, xml.encode("utf-8"), doc.part.package)
    doc.part.relate_to(part, rt)


def _member_xml(docx_path: str, member: str) -> str:
    """Raw, bytes-exact XML string of a zip member (NBSP-sensitive substring checks)."""
    with zipfile.ZipFile(docx_path) as z:
        return z.read(member).decode("utf-8")


def _no_tracked_elements(xml: str) -> bool:
    """True iff the parsed part has ZERO w:ins, w:del and w:delText elements. Element-level
    (not substring) so it never false-matches on unrelated tags like <w:instrText>."""
    root = parse_xml(xml.encode("utf-8"))
    return (
        not root.findall(".//" + qn("w:ins"))
        and not root.findall(".//" + qn("w:del"))
        and not root.findall(".//" + qn("w:delText"))
    )


# --------------------------------------------------------------------------- body ins/del

def _build_body_fixture(path) -> str:
    doc = Document()
    p = doc.add_paragraph("Normalny text pred zmenou. ")
    p._p.append(parse_xml(_ins_xml(INS_MARK)))
    p._p.append(parse_xml(_del_xml(DEL_MARK)))
    doc.save(str(path))
    return str(path)


def test_body_tracked_changes_accepted(tmp_path):
    in_path = _build_body_fixture(tmp_path / "in.docx")
    # Sanity: the raw fixture really carries both markers + a w:del, so a later "gone" assertion
    # is a genuine RED->GREEN signal, not a vacuous pass.
    raw_in = _member_xml(in_path, "word/document.xml")
    assert DEL_MARK in raw_in and INS_MARK in raw_in
    assert "<w:del " in raw_in

    out_path = str(tmp_path / "out.docx")
    redact_docx_body(in_path, out_path)
    xml = _member_xml(out_path, "word/document.xml")

    # deleted subtree physically gone
    assert DEL_MARK not in xml, f"LEAKED deleted text {DEL_MARK!r} survived in {xml!r}"
    assert _no_tracked_elements(xml), f"tracked-change element survived in {xml!r}"
    # inserted content retained as ordinary body text (inside a <w:t>, no longer wrapped by ins)
    assert INS_MARK in xml
    doc = Document(out_path)
    assert INS_MARK in doc.paragraphs[0].text
    # and it lands in a real <w:t> run, not orphaned
    root = parse_xml(xml.encode("utf-8"))
    ts = [t.text for t in root.findall(".//" + qn("w:t")) if t.text and INS_MARK in t.text]
    assert ts, "INSERTEDKEEP not found inside any <w:t>"


# --------------------------------------------------------------------------- note parts

def _build_notes_fixture(path) -> str:
    """Docx with a comment (element path) and a footnote (blob path), each carrying a <w:del>."""
    doc = Document()
    doc.add_paragraph("Telo bez PII.")

    cm_body = (
        f'<w:comment w:id="1" w:author="Advokat" w:date="2024-01-01T00:00:00Z" '
        f'w:initials="AK"><w:p><w:r><w:t xml:space="preserve">Komentar. </w:t></w:r>'
        f"{_del_xml('DELETEDCOMMENT', wid='7001', ns=False)}</w:p></w:comment>"
    )
    fn_body = (
        f'<w:footnote w:id="2"><w:p><w:r><w:t xml:space="preserve">Poznamka. </w:t></w:r>'
        f"{_del_xml('DELETEDFOOTNOTE', wid='7002', ns=False)}</w:p></w:footnote>"
    )

    _attach(doc, "comments.xml", _CM_CT, f'<w:comments {nsdecls("w")}>{cm_body}</w:comments>', RT.COMMENTS)
    _attach(doc, "footnotes.xml", _FN_CT, f'<w:footnotes {nsdecls("w")}>{_FN_SEPS}{fn_body}</w:footnotes>', RT.FOOTNOTES)
    doc.save(str(path))
    return str(path)


def test_note_part_tracked_del_stripped_both_paths(tmp_path):
    in_path = _build_notes_fixture(tmp_path / "in.docx")
    assert "DELETEDCOMMENT" in _member_xml(in_path, "word/comments.xml")
    assert "DELETEDFOOTNOTE" in _member_xml(in_path, "word/footnotes.xml")

    out_path = str(tmp_path / "out.docx")
    redact_docx_body(in_path, out_path)

    cm = _member_xml(out_path, "word/comments.xml")   # element (CommentsPart) path
    assert "DELETEDCOMMENT" not in cm, f"LEAKED in comments: {cm!r}"
    assert _no_tracked_elements(cm)
    assert "Komentar." in cm  # ordinary comment text untouched

    fn = _member_xml(out_path, "word/footnotes.xml")  # blob (generic Part) path
    assert "DELETEDFOOTNOTE" not in fn, f"LEAKED in footnotes: {fn!r}"
    assert _no_tracked_elements(fn)
    assert "<w:separator/>" in fn and "<w:continuationSeparator/>" in fn  # seps survive
    Document(out_path)  # reopens without corruption


# --------------------------------------------------------------------------- insertion still redacts

def _build_ins_pii_fixture(path) -> str:
    doc = Document()
    p = doc.add_paragraph("Subjekt eviduje ICO ")
    p._p.append(parse_xml(_ins_xml(ICO, wid="9201")))
    doc.save(str(path))
    return str(path)


def test_inserted_pii_flows_through_redaction(tmp_path):
    """An insertion that CONTAINS auto-detectable PII must, once unwrapped, still be redacted by
    the normal W1 pass — proving the strip runs BEFORE redaction and does not special-case
    insertions. This only holds if ordering is correct."""
    in_path = _build_ins_pii_fixture(tmp_path / "in.docx")
    assert ICO in _member_xml(in_path, "word/document.xml")

    out_path = str(tmp_path / "out.docx")
    redact_docx_body(in_path, out_path)
    xml = _member_xml(out_path, "word/document.xml")

    assert ICO not in xml, f"inserted PII {ICO!r} leaked un-redacted in {xml!r}"
    assert "[ICO_1]" in xml
    assert _no_tracked_elements(xml)
    Document(out_path)

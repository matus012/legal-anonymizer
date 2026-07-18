"""W4b (context.md §10): scrub document metadata that carries PII — three targets, three
storage shapes:

* docProps/core.xml — exposed via python-docx's ``doc.core_properties`` API (dc:creator,
  cp:lastModifiedBy, and the other same-risk str properties).
* docProps/app.xml — python-docx has NO API for this; it is a generic blob-backed Part.
  Company/Manager are PII-bearing free-text fields inside it.
* word/comments.xml ``<w:comment w:author=...>`` — deferred from W3 (see
  test_writer_docx_notes.py's ``test_comment_element_path_persists_through_save``, which
  explicitly notes "author attribute is W4 scope — W3 leaves it").

These are BLANKED UNCONDITIONALLY BY POSITION — detect() never runs over metadata, and the
original value is never preserved, so there is no "does this look like PII" branch to test:
every fixture value must be gone in the output regardless of its content.

Fixtures are built entirely in-memory (this module must NOT import corpus/): a minimal
python-docx Document, its own default-created app.xml Part (python-docx creates one even for
a blank Document — verified by probe), plus a hand-attached comments part (same pattern as
test_writer_docx_notes.py). redact_docx_body takes file paths, so every test writes the
fixture to disk first and redacts path->path exactly as production does.

RED/GREEN: run this file against the pre-W4b writer and every marker-presence assertion FAILS
because core.xml/app.xml/comments.xml still carry the marker untouched. After wiring the
metadata scrub into redact_docx_body they pass: markers gone, unrelated content byte-preserved.
"""
from __future__ import annotations

import zipfile

from docx import Document
from docx.opc.constants import RELATIONSHIP_TYPE as RT
from docx.opc.packuri import PackURI
from docx.opc.part import Part
from docx.oxml.ns import nsdecls

from writer.docx_body import redact_docx_body

_CM_CT = "application/vnd.openxmlformats-officedocument.wordprocessingml.comments+xml"

CORE_MARKER = "METACREATOR"
COMPANY_MARKER = "METACOMPANY"
MANAGER_MARKER = "METAMANAGER"
AUTHOR_MARKER = "METAAUTHOR"
INITIALS_MARKER = "MA"

# Hand-crafted app.xml blob: two PII-bearing markers (Company, Manager) plus unrelated tags
# (Application, a HeadingPairs vt:vector) that must survive byte-identical, proving a
# surgical scrub rather than wholesale part deletion/reserialization.
_APP_XML = (
    "<?xml version=\"1.0\" encoding=\"UTF-8\" standalone=\"yes\"?>\n"
    '<Properties xmlns="http://schemas.openxmlformats.org/officeDocument/2006/extended-properties" '
    'xmlns:vt="http://schemas.openxmlformats.org/officeDocument/2006/docPropsVTypes">'
    "<Application>Word</Application>"
    "<HeadingPairs><vt:vector size=\"2\" baseType=\"variant\">"
    "<vt:variant><vt:lpstr>Title</vt:lpstr></vt:variant>"
    "<vt:variant><vt:i4>1</vt:i4></vt:variant></vt:vector></HeadingPairs>"
    f"<Company>{COMPANY_MARKER}</Company>"
    f"<Manager>{MANAGER_MARKER}</Manager>"
    "</Properties>"
)


def _app_part(doc):
    for part in doc.part.package.iter_parts():
        if str(part.partname) == "/docProps/app.xml":
            return part
    raise AssertionError("python-docx did not create a default /docProps/app.xml part")


def _attach_comments(doc) -> None:
    cm_body = (
        f'<w:comment w:id="1" w:author="{AUTHOR_MARKER}" w:date="2024-01-01T00:00:00Z" '
        f'w:initials="{INITIALS_MARKER}"><w:p><w:r>'
        '<w:t xml:space="preserve">Bezne znenie komentara bez PII.</w:t>'
        "</w:r></w:p></w:comment>"
    )
    part = Part(
        PackURI("/word/comments.xml"),
        _CM_CT,
        f'<w:comments {nsdecls("w")}>{cm_body}</w:comments>'.encode("utf-8"),
        doc.part.package,
    )
    doc.part.relate_to(part, RT.COMMENTS)


def _build_fixture(path) -> str:
    doc = Document()
    doc.add_paragraph("Telo dokumentu bez PII.")

    cp = doc.core_properties
    cp.author = CORE_MARKER
    cp.last_modified_by = CORE_MARKER
    cp.title = CORE_MARKER
    cp.subject = CORE_MARKER
    cp.keywords = CORE_MARKER
    cp.category = CORE_MARKER
    cp.comments = CORE_MARKER

    app_part = _app_part(doc)
    app_part._blob = _APP_XML.encode("utf-8")

    _attach_comments(doc)

    doc.save(str(path))
    return str(path)


def _part_xml(docx_path: str, member: str) -> str:
    """Raw, bytes-exact XML string of a zip member (NBSP-sensitive substring checks)."""
    with zipfile.ZipFile(docx_path) as z:
        return z.read(member).decode("utf-8")


def _redacted(tmp_path) -> str:
    in_path = _build_fixture(tmp_path / "in.docx")
    # Pre-change sanity: the raw fixture really does contain every target, so a later
    # "target absent" assertion is a genuine RED->GREEN signal, not a vacuous pass.
    core_in = _part_xml(in_path, "docProps/core.xml")
    app_in = _part_xml(in_path, "docProps/app.xml")
    comments_in = _part_xml(in_path, "word/comments.xml")
    assert CORE_MARKER in core_in
    assert COMPANY_MARKER in app_in and MANAGER_MARKER in app_in
    assert f'w:author="{AUTHOR_MARKER}"' in comments_in

    out_path = str(tmp_path / "out.docx")
    redact_docx_body(in_path, out_path)
    return out_path


def test_core_properties_scrubbed(tmp_path):
    out = _redacted(tmp_path)
    xml = _part_xml(out, "docProps/core.xml")
    assert CORE_MARKER not in xml, f"LEAKED core.xml marker: {CORE_MARKER!r} survived in {xml!r}"
    assert "<dc:creator/>" in xml or "<dc:creator></dc:creator>" in xml
    assert "<cp:lastModifiedBy/>" in xml or "<cp:lastModifiedBy></cp:lastModifiedBy>" in xml
    # dates are NOT PII -- must survive untouched.
    assert "dcterms:created" in xml and "dcterms:modified" in xml
    Document(out)


def test_app_xml_company_manager_scrubbed(tmp_path):
    out = _redacted(tmp_path)
    xml = _part_xml(out, "docProps/app.xml")
    assert COMPANY_MARKER not in xml, f"LEAKED app.xml Company: {COMPANY_MARKER!r} survived in {xml!r}"
    assert MANAGER_MARKER not in xml, f"LEAKED app.xml Manager: {MANAGER_MARKER!r} survived in {xml!r}"
    assert "<Company></Company>" in xml or "<Company/>" in xml
    assert "<Manager></Manager>" in xml or "<Manager/>" in xml
    # unrelated tags byte-preserved -- proves a surgical scrub, not wholesale reserialization.
    assert "<Application>Word</Application>" in xml
    assert '<vt:vector size="2" baseType="variant">' in xml
    assert xml.startswith('<?xml version="1.0" encoding="UTF-8" standalone="yes"?>')
    Document(out)


def test_comment_author_and_initials_scrubbed(tmp_path):
    out = _redacted(tmp_path)
    xml = _part_xml(out, "word/comments.xml")
    assert AUTHOR_MARKER not in xml, f"LEAKED comment author: {AUTHOR_MARKER!r} survived in {xml!r}"
    assert INITIALS_MARKER not in xml, f"LEAKED comment initials: {INITIALS_MARKER!r} survived in {xml!r}"
    assert 'w:author=""' in xml
    assert 'w:initials=""' in xml
    # id/date preserved, comment body text untouched.
    assert 'w:id="1"' in xml
    assert 'w:date="2024-01-01T00:00:00Z"' in xml
    assert "Bezne znenie komentara bez PII." in xml
    Document(out)


def test_comment_scrub_persists_through_save(tmp_path):
    """Comments reopen as a CommentsPart whose .blob RE-SERIALIZES from .element (see
    test_writer_docx_notes.py), so the author scrub must mutate part.element in place; a
    ._blob reassignment would be silently discarded. Reading the SAVED output zip (not the
    in-memory tree) proves the element mutation actually persisted through doc.save()."""
    out = _redacted(tmp_path)
    xml = _part_xml(out, "word/comments.xml")
    assert AUTHOR_MARKER not in xml
    doc = Document(out)
    (comments_part,) = [
        r.target_part for r in doc.part.rels.values() if r.reltype.endswith("comments")
    ]
    reopened = comments_part.blob.decode("utf-8")
    assert AUTHOR_MARKER not in reopened
    assert 'w:author=""' in reopened


def test_body_still_clean_and_all_parts_reopen(tmp_path):
    """End-to-end: after W4b the whole document still reopens and the non-metadata body is
    intact."""
    out = _redacted(tmp_path)
    doc = Document(out)
    assert doc.paragraphs[0].text == "Telo dokumentu bez PII."
    for member in ("docProps/core.xml", "docProps/app.xml", "word/comments.xml"):
        _part_xml(out, member)  # member exists / readable

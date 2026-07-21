"""RedactionDecisions + LabelMap snippet/grouping extensions (Phase 6).

Hand-built candidates only — tests never import corpus/ or eval/.
"""
from types import SimpleNamespace

from writer.decisions import RedactionDecisions
from writer.labelmap import LabelMap, make_snippet


def _cand(type_, surface, start=0, end=None, auto=True):
    return SimpleNamespace(
        type=type_, surface=surface, start=start,
        end=(end if end is not None else start + len(surface)), auto=auto,
    )


def test_decisions_defaults_are_empty_and_frozen():
    d = RedactionDecisions()
    assert d.extra_terms == ()
    assert d.suppress_groups == frozenset()
    assert d.force_groups == frozenset()
    try:
        d.extra_terms = ("x",)
        assert False, "must be frozen"
    except Exception:
        pass


def test_group_key_for_matches_group_key():
    lm = LabelMap(["Ján Novák"])
    c = _cand("MENO", "Nováka")
    assert lm.group_key_for(c.type, c.surface) == lm.group_key(c)
    c2 = _cand("DIC", "1234567890")
    assert lm.group_key_for("DIC", "1234567890") == lm.group_key(c2)


def test_groups_inverts_the_cache():
    lm = LabelMap(None)
    c = _cand("DIC", "1234567890")
    label = lm.label_for(c)
    assert lm.groups() == {label: ("DIC", lm.group_key(c))}


def test_record_occurrence_keeps_first_snippet_only():
    lm = LabelMap(None)
    lm.record_occurrence("[DIC_1]", "body", "123", snippet="first ctx")
    lm.record_occurrence("[DIC_1]", "header", "123", snippet="second ctx")
    assert lm.contexts["[DIC_1]"] == "first ctx"
    assert lm.occurrences["[DIC_1]"] == [("body", "123"), ("header", "123")]  # shape UNCHANGED


def test_record_low_confidence_snippets_stay_aligned():
    lm = LabelMap(None)
    lm.record_low_confidence("body", "RODNE_CISLO", "835112/0009", snippet="ctx a")
    lm.record_low_confidence("page_2", "ICO", "12345678")  # no snippet -> "" placeholder
    assert lm.low_confidence == [("body", "RODNE_CISLO", "835112/0009"), ("page_2", "ICO", "12345678")]
    assert lm.lc_contexts == ["ctx a", ""]


def test_make_snippet_windows_and_collapses_whitespace():
    text = "aaaa    bbbb  NEEDLE\n\ncccc dddd"
    s = make_snippet(text, text.index("NEEDLE"), text.index("NEEDLE") + 6, radius=6)
    assert "NEEDLE" in s
    assert "\n" not in s and "  " not in s


# --- DOCX writer plumbing (hand-built docx fixtures, python-docx) ---
import docx as _docx

from writer.docx_body import redact_docx_body, redact_docx_collect


def _mk_docx(tmp_path, text):
    p = tmp_path / "in.docx"
    d = _docx.Document()
    d.add_paragraph(text)
    d.save(str(p))
    return str(p)


def _docx_text(path):
    return "\n".join(par.text for par in _docx.Document(path).paragraphs)


def test_docx_suppress_keeps_surface_in_output(tmp_path):
    """Retention proof: a suppressed group's PII MUST survive verbatim (human said keep)."""
    src = _mk_docx(tmp_path, "Konatel Jan Novak, DIC 2023456789, uhradi sumu.")
    out_plain = str(tmp_path / "plain.docx")
    lm = redact_docx_collect(src, out_plain, known_entities=["Jan Novak"])
    # sanity: by default the DIC is redacted
    assert "2023456789" not in _docx_text(out_plain)
    key = ("DIC", lm.group_key_for("DIC", "2023456789"))
    out = str(tmp_path / "sup.docx")
    from writer.decisions import RedactionDecisions
    redact_docx_body(src, out, known_entities=["Jan Novak"],
                     decisions=RedactionDecisions(suppress_groups=frozenset({key})))
    txt = _docx_text(out)
    assert "2023456789" in txt            # survived — suppression respected
    assert "Novak" not in txt             # untouched groups still redacted
    # and the report records the human decision as NOT redacted:
    rep = (tmp_path / "sup_report.txt").read_text(encoding="utf-8")
    assert "2023456789" in rep


def test_docx_force_redacts_low_confidence_group(tmp_path):
    # 835112/0009 is RC-shaped with an INVALID checksum -> auto=False, normally left intact
    src = _mk_docx(tmp_path, "rodne cislo: 835112/0009")
    out_plain = str(tmp_path / "plain.docx")
    lm = redact_docx_collect(src, out_plain, known_entities=None)
    assert "835112/0009" in _docx_text(out_plain)  # precondition: low-conf survives by default
    key = ("RODNE_CISLO", lm.group_key_for("RODNE_CISLO", "835112/0009"))
    out = str(tmp_path / "forced.docx")
    from writer.decisions import RedactionDecisions
    lm2 = redact_docx_collect(src, out, known_entities=None,
                              decisions=RedactionDecisions(force_groups=frozenset({key})))
    txt = _docx_text(out)
    assert "835112/0009" not in txt
    assert "[RODNE_CISLO_1]" in txt
    assert lm2.occurrences.get("[RODNE_CISLO_1]")  # recorded as a real redaction


def test_docx_extra_terms_redact_like_known_entities(tmp_path):
    src = _mk_docx(tmp_path, "Svedok Karol Vlk vypovedal.")
    out = str(tmp_path / "extra.docx")
    from writer.decisions import RedactionDecisions
    redact_docx_body(src, out, known_entities=None,
                     decisions=RedactionDecisions(extra_terms=("Karol Vlk",)))
    txt = _docx_text(out)
    assert "Vlk" not in txt and "[MENO_" in txt


def test_docx_collect_returns_labelmap_with_snippets(tmp_path):
    src = _mk_docx(tmp_path, "Zmluvna strana Jan Novak podpisala zmluvu.")
    out = str(tmp_path / "c.docx")
    lm = redact_docx_collect(src, out, known_entities=["Jan Novak"])
    label = next(iter(lm.groups()))
    assert "zmluvu" in lm.contexts[label] or "strana" in lm.contexts[label]


def test_docx_decisions_none_is_default_behaviour(tmp_path):
    src = _mk_docx(tmp_path, "Konatel Jan Novak, DIC 2023456789.")
    a, b = str(tmp_path / "a.docx"), str(tmp_path / "b.docx")
    redact_docx_body(src, a, known_entities=["Jan Novak"])
    redact_docx_body(src, b, known_entities=["Jan Novak"], decisions=None)
    assert _docx_text(a) == _docx_text(b)

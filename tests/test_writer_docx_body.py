"""W1 unit tests (context.md §10): redact PII in DOCX top-level body paragraphs, remapping
detect() spans onto <w:r> runs even when a surface starts and/or ends MID-RUN.

Fixtures are hand-built in memory with python-docx (NO corpus import). Each test first
asserts the naive/pre-fix precondition (the surface is still extractable, or an rPr-bearing
run is intact), then asserts redact_docx_body destroyed the surface while preserving the
surrounding text and formatting. The input file is never mutated.
"""
from __future__ import annotations

from docx import Document
from docx.oxml import parse_xml
from docx.oxml.ns import qn
from docx.text.paragraph import Paragraph

from writer.docx_body import redact_docx_body


def _build(paragraphs: list[list[tuple]], path) -> None:
    """paragraphs: list of paragraphs; each paragraph a list of (text, bold?) run specs.

    Runs are emitted verbatim and in order, so a caller can split a single logical surface
    across several <w:r> runs (including mid-token) exactly like the real corpus does.
    """
    doc = Document()
    for runs in paragraphs:
        p = doc.add_paragraph()
        for spec in runs:
            text = spec[0]
            bold = spec[1] if len(spec) > 1 else None
            r = p.add_run(text)
            if bold:
                r.bold = True
    doc.save(str(path))


def _body_text(path) -> str:
    """Top-level body reconstruction, identical to the leak-gate probe: join paragraph.text
    (== concatenation of run.text) with NO separator inside a paragraph (context.md §10)."""
    doc = Document(str(path))
    return "\n".join(p.text for p in doc.paragraphs)


# ---------------------------------------------------------------- (a) mid-token, run-aligned
def test_name_split_mid_token_across_three_runs_is_redacted(tmp_path):
    src, out = tmp_path / "a_in.docx", tmp_path / "a_out.docx"
    # runs reconstruct to 'Lucie Molnárovej' (declined) — the corpus split-run shape.
    _build([[("Luc",), ("ie Molnáro",), ("vej",)]], src)

    # precondition: the glued surface is extractable from the input.
    assert "Lucie Molnárovej" in _body_text(src)

    redact_docx_body(str(src), str(out), known_entities=["Lucia Molnárová"])

    txt = _body_text(out)
    assert "Lucie" not in txt
    assert "Molnárovej" not in txt
    assert "[MENO]" in txt
    # input file untouched.
    assert "Lucie Molnárovej" in _body_text(src)


# ---------------------------------------------------------------- (b) starts AND ends mid-run
def test_name_starting_and_ending_mid_run_is_redacted_keeping_surrounding_text(tmp_path):
    src, out = tmp_path / "b_in.docx", tmp_path / "b_out.docx"
    # 'Jánom Novákom' starts inside run 0 (after 'Zmluva medzi ') and ends inside run 1
    # (before ' dnes') — the case the corpus never produces and W1 exists to handle.
    _build([[("Zmluva medzi Ján",), ("om Novákom dnes",)]], src)

    assert "Jánom Novákom" in _body_text(src)

    redact_docx_body(str(src), str(out), known_entities=["Ján Novák"])

    txt = _body_text(out)
    assert "Jánom" not in txt
    assert "Novákom" not in txt
    assert "[MENO]" in txt
    # the non-PII head and tail glued into the boundary runs must survive intact.
    assert txt.startswith("Zmluva medzi ")
    assert txt.endswith(" dnes")


# ---------------------------------------------------------------- (c) rPr survives a partial hit
def test_partially_overlapped_bold_run_keeps_its_rpr_on_surviving_fragment(tmp_path):
    src, out = tmp_path / "c_in.docx", tmp_path / "c_out.docx"
    # run 1 is bold and holds 'Ján Novák súhlasí'; only 'Ján Novák' is redacted, so the
    # surviving ' súhlasí' fragment must stay bold.
    _build([[("Kontakt: ",), ("Ján Novák súhlasí", True)]], src)

    # precondition: bold run is intact and the surface is extractable.
    src_doc = Document(str(src))
    src_bold = [r for p in src_doc.paragraphs for r in p.runs if r.bold]
    assert any("Ján Novák" in r.text for r in src_bold)

    redact_docx_body(str(src), str(out), known_entities=["Ján Novák"])

    txt = _body_text(out)
    assert "Ján Novák" not in txt
    assert "[MENO]" in txt

    doc = Document(str(out))
    survivors = [r for p in doc.paragraphs for r in p.runs if "súhlasí" in r.text]
    assert survivors, "surviving ' súhlasí' fragment vanished"
    assert all(r.bold for r in survivors), "surviving fragment lost its <w:rPr> (bold)"


# ---------------------------------------------------------------- (d) self-detecting identifier
def test_valid_ico_self_detects_without_known_entities(tmp_path):
    src, out = tmp_path / "d_in.docx", tmp_path / "d_out.docx"
    # 12345679 is a checksum-valid IČO (auto=True), matched with no known_entities.
    _build([[("IČO spoločnosti 12345679 tu.",)]], src)

    assert "12345679" in _body_text(src)

    redact_docx_body(str(src), str(out))  # known_entities omitted

    txt = _body_text(out)
    assert "12345679" not in txt
    assert "[ICO]" in txt


# ---------------------------------------------------------------- output integrity
def test_output_reopens_cleanly_and_input_is_never_mutated(tmp_path):
    src, out = tmp_path / "e_in.docx", tmp_path / "e_out.docx"
    _build([[("Kupujúci ",), ("Ján Novák",), (" a IČO 12345679.",)]], src)
    before = src.read_bytes()

    redact_docx_body(str(src), str(out), known_entities=["Ján Novák"])

    # reopens without XML corruption
    reopened = Document(str(out))
    assert reopened.paragraphs, "output has no body paragraphs"
    # input bytes unchanged
    assert src.read_bytes() == before


# ===================================================================== W2 (context.md §10)
# W2 extends coverage from doc.paragraphs to the OTHER <w:p> locations Word hides PII in:
# table cells, headers, footers, header/footer tables, and VML textboxes (body + hdr/ftr).
# The SAME run-remap core (_redact_paragraph) is reused verbatim; these tests only prove the
# extra locations are now walked. Each first asserts the surface is extractable (RED before
# W2 wiring exists) then asserts redaction. Fixtures are hand-built (NO corpus import).

# VML textbox: python-docx has no paragraph API for it, so it is built as raw <w:pict> XML
# and appended to a real body run, exactly like the corpus <v:textbox> shape.
_PICT_TEXTBOX = (
    '<w:pict xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" '
    'xmlns:v="urn:schemas-microsoft-com:vml">'
    "<v:shape><v:textbox><w:txbxContent>"
    '<w:p><w:r><w:t xml:space="preserve">{inner}</w:t></w:r></w:p>'
    "</w:txbxContent></v:textbox></v:shape></w:pict>"
)


def _table_cell_texts(path) -> list[str]:
    """Every table-cell paragraph's text, dedup'd by <w:tc> identity (merged cells share a
    tc), matching the leak-gate probe."""
    doc = Document(str(path))
    out: list[str] = []
    seen: set[int] = set()
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if id(cell._tc) in seen:
                    continue
                seen.add(id(cell._tc))
                out.extend(p.text for p in cell.paragraphs)
    return out


def _header_footer_texts(path) -> list[str]:
    doc = Document(str(path))
    out: list[str] = []
    for section in doc.sections:
        for hf in (section.header, section.footer):
            out.extend(p.text for p in hf.paragraphs)
    return out


def _textbox_texts(path) -> list[str]:
    """Text of every <w:p> inside every <w:txbxContent> in body + header + footer parts."""
    doc = Document(str(path))
    out: list[str] = []
    elems = [doc.element.body]
    for section in doc.sections:
        elems.append(section.header._element)
        elems.append(section.footer._element)
    for el in elems:
        for tb in el.findall(".//" + qn("w:txbxContent")):
            for p_elem in tb.findall(qn("w:p")):
                out.append(Paragraph(p_elem, doc).text)
    return out


# ------------------------------------------------------------- (t1) table cell split 3 runs
def test_table_cell_name_split_across_three_runs_is_redacted(tmp_path):
    src, out = tmp_path / "t1_in.docx", tmp_path / "t1_out.docx"
    doc = Document()
    table = doc.add_table(rows=1, cols=1)
    cell_p = table.cell(0, 0).paragraphs[0]
    # 'Lucie Molnárovej' split mid-token across 3 runs, exactly like the corpus 'bod 7.6'.
    for frag in ("Luc", "ie Molnáro", "vej"):
        cell_p.add_run(frag)
    doc.save(str(src))

    # precondition: the glued surface is extractable from a table cell.
    assert any("Lucie Molnárovej" in t for t in _table_cell_texts(src))

    redact_docx_body(str(src), str(out), known_entities=["Lucia Molnárová"])

    cells = _table_cell_texts(out)
    joined = "".join(cells)
    assert "Lucie" not in joined
    assert "Molnárovej" not in joined
    assert any("[MENO]" in t for t in cells)


# ---------------------------------------------------------- (t2) merged cell processed once
def test_merged_table_cell_is_processed_exactly_once(tmp_path):
    # A vertically merged cell makes table.cell(0,0) and table.cell(1,0) return the SAME
    # <w:tc>; without dedup its paragraph would be walked twice. The observable contract is
    # that the merged cell is redacted EXACTLY like an equivalent single cell — same labels,
    # not a doubled rewrite — and the file stays valid. (detect emits one [MENO] per declined
    # name token, so "Ján Novák" itself yields two labels; the control cell captures that so
    # the test asserts the merged path matches it rather than hard-coding detect's count.)
    control_src, control_out = tmp_path / "t2c_in.docx", tmp_path / "t2c_out.docx"
    cdoc = Document()
    cdoc.add_table(rows=1, cols=1).cell(0, 0).paragraphs[0].add_run("Ján Novák")
    cdoc.save(str(control_src))
    redact_docx_body(str(control_src), str(control_out), known_entities=["Ján Novák"])
    expected_labels = "".join(_table_cell_texts(control_out)).count("[MENO]")

    src, out = tmp_path / "t2_in.docx", tmp_path / "t2_out.docx"
    doc = Document()
    table = doc.add_table(rows=2, cols=1)
    table.cell(0, 0).paragraphs[0].add_run("Ján Novák")
    table.cell(0, 0).merge(table.cell(1, 0))  # cell(0,0) & cell(1,0) now share one <w:tc>
    doc.save(str(src))

    # sanity: the two logical cells really are one tc (so a naive per-cell walk double-visits).
    src_doc = Document(str(src))
    assert id(src_doc.tables[0].cell(0, 0)._tc) == id(src_doc.tables[0].cell(1, 0)._tc)
    assert any("Ján Novák" in t for t in _table_cell_texts(src))

    redact_docx_body(str(src), str(out), known_entities=["Ján Novák"])

    joined = "".join(_table_cell_texts(out))
    assert "Ján Novák" not in joined
    # processed exactly once: same label count as the single-cell control, not doubled.
    assert joined.count("[MENO]") == expected_labels
    # and the output re-opens without corruption.
    reopened = Document(str(out))
    assert reopened.tables, "output lost its table"


# ------------------------------------------------------- (t3) header: email + IČO both gone
def test_header_paragraph_redacts_email_and_ico(tmp_path):
    src, out = tmp_path / "t3_in.docx", tmp_path / "t3_out.docx"
    doc = Document()
    header = doc.sections[0].header
    header.is_linked_to_previous = False
    # one run carrying TWO disjoint detect spans (EMAIL + a checksum-valid IČO).
    header.paragraphs[0].add_run("Kontakt a@b.sk alebo IČO 12345679 tu")
    doc.save(str(src))

    hf = _header_footer_texts(src)
    assert any("a@b.sk" in t and "12345679" in t for t in hf)

    redact_docx_body(str(src), str(out), known_entities=None)

    hf = _header_footer_texts(out)
    joined = "".join(hf)
    assert "a@b.sk" not in joined
    assert "12345679" not in joined
    assert "[EMAIL]" in joined
    assert "[ICO]" in joined


# --------------------------------------------------------------- (t4) footer: name redacted
def test_footer_paragraph_name_is_redacted(tmp_path):
    src, out = tmp_path / "t4_in.docx", tmp_path / "t4_out.docx"
    doc = Document()
    footer = doc.sections[0].footer
    footer.is_linked_to_previous = False
    footer.paragraphs[0].add_run("Ján Novák")
    doc.save(str(src))

    assert any("Ján Novák" in t for t in _header_footer_texts(src))

    redact_docx_body(str(src), str(out), known_entities=["Ján Novák"])

    hf = _header_footer_texts(out)
    joined = "".join(hf)
    assert "Ján Novák" not in joined
    assert "[MENO]" in joined


# ------------------------------------------------------ (t5) VML textbox: detectable id gone
def test_vml_textbox_detectable_identifier_is_redacted(tmp_path):
    src, out = tmp_path / "t5_in.docx", tmp_path / "t5_out.docx"
    doc = Document()
    p = doc.add_paragraph()
    r = p.add_run()
    # inner textbox paragraph holds a checksum-valid DIČ (a DETECTABLE type, not a gazetteer
    # name) so redaction depends on detect(), not on known_entities.
    r._r.append(parse_xml(_PICT_TEXTBOX.format(inner="Danova cislo DIC 2020123456 tu")))
    doc.save(str(src))

    # precondition: the DIČ surface is extractable from the textbox (and NOT from the body,
    # since the outer run carries no <w:t>).
    assert any("2020123456" in t for t in _textbox_texts(src))

    redact_docx_body(str(src), str(out), known_entities=None)

    tb = _textbox_texts(out)
    joined = "".join(tb)
    assert "2020123456" not in joined
    assert "[DIC]" in joined

"""Phase 6 Task 4: gui.model — pure scan/decisions/export logic, no Qt.

Hand-built fixtures via the same helpers as tests/test_writer_decisions.py and
tests/test_writer_decisions_pdf.py — tests never import corpus/ or eval/.
"""
import docx as _docx
import fitz

from gui.model import FileScan, ReviewRow, build_decisions, export_file, out_path_for, scan_file


def _mk_docx(tmp_path, text):
    p = tmp_path / "in.docx"
    d = _docx.Document()
    d.add_paragraph(text)
    d.save(str(p))
    return str(p)


def _docx_text(path):
    return "\n".join(par.text for par in _docx.Document(path).paragraphs)


def _mk_pdf(tmp_path, text):
    p = tmp_path / "in.pdf"
    doc = fitz.open()
    page = doc.new_page()
    page.insert_text((72, 72), text)
    doc.save(str(p))
    doc.close()
    return str(p)


def _pdf_text(path):
    doc = fitz.open(path)
    t = "\n".join(pg.get_text("text") for pg in doc)
    doc.close()
    return t


def test_out_path_for_appends_anon(tmp_path):
    assert out_path_for(r"C:\x\zmluva.docx").endswith("zmluva_anon.docx")
    assert out_path_for(r"C:\x\zaloba.pdf").endswith("zaloba_anon.pdf")


def test_scan_docx_builds_grouped_rows(tmp_path):
    src = _mk_docx(tmp_path, "Jan Novak a Novakovi patri DIC 2023456789. rodne cislo 835112/0009.")
    scan = scan_file(src, ["Jan Novak"])
    assert scan.error is None
    auto = [r for r in scan.rows if r.bucket == "auto"]
    review = [r for r in scan.rows if r.bucket == "review"]
    meno = next(r for r in auto if r.type == "MENO")
    # detect.known_entities emits one MENO per matched TOKEN: Jan + Novak + Novakovi ->
    # three occurrences, all bound to ONE group (entity 0) by the declension matcher.
    assert meno.count == 3
    assert meno.locations == ("body",)
    assert any(r.type == "DIC" for r in auto)
    assert any(r.type == "RODNE_CISLO" for r in review)
    assert all(r.snippet for r in scan.rows)    # every row carries context


def test_scan_leaves_no_temp_files(tmp_path):
    src = _mk_docx(tmp_path, "DIC 2023456789")
    scan_file(src, None)
    assert list(tmp_path.iterdir()) == [tmp_path / "in.docx"]  # nothing new beside the source


def test_scan_pdf_without_text_layer_reports_error(tmp_path):
    p = tmp_path / "scan.pdf"
    doc = fitz.open(); doc.new_page(); doc.save(str(p)); doc.close()
    scan = scan_file(str(p), None)
    assert scan.error is not None and scan.rows == []


def test_build_decisions_from_checkbox_state(tmp_path):
    src = _mk_docx(tmp_path, "Jan Novak, DIC 2023456789, rodne cislo 835112/0009.")
    scan = scan_file(src, ["Jan Novak"])
    dic = next(r for r in scan.rows if r.type == "DIC")
    rc = next(r for r in scan.rows if r.type == "RODNE_CISLO")
    checked = {r.group: (r.bucket == "auto") for r in scan.rows}
    checked[dic.group] = False    # human unticks the DIC
    checked[rc.group] = True      # human ticks the low-confidence RC
    d = build_decisions(scan.rows, checked, extra_terms=("Karol Vlk",))
    assert dic.group in d.suppress_groups
    assert rc.group in d.force_groups
    assert d.extra_terms == ("Karol Vlk",)
    meno = next(r for r in scan.rows if r.type == "MENO")
    assert meno.group not in d.suppress_groups


def test_export_writes_anon_and_report(tmp_path):
    src = _mk_docx(tmp_path, "Jan Novak, DIC 2023456789.")
    scan = scan_file(src, ["Jan Novak"])
    checked = {r.group: (r.bucket == "auto") for r in scan.rows}
    out, report = export_file(src, ["Jan Novak"], build_decisions(scan.rows, checked, ()))
    assert out.endswith("in_anon.docx") and report.endswith("in_anon_report.txt")
    txt = "\n".join(p.text for p in _docx.Document(out).paragraphs)
    assert "Novak" not in txt and "2023456789" not in txt

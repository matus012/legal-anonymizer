"""Regression tests for the table-cell dedup key in writer.docx_body._redact_cells.

Bug (fixed): _redact_cells deduped cells with ``if id(cell._tc) in seen``. row.cells mints a
fresh _Cell proxy per row and each row's proxies are GC'd before the next row's are allocated,
so CPython reuses the freed address and a later DISTINCT cell's proxy collides with an earlier
cell's id() in ``seen`` -> the later cell is silently skipped and its PII leaks. The fix keys
the set on the lxml <w:tc> element object (stable for the table's lifetime, identity-hashed)
instead of id() of the transient proxy.

These fixtures build docx tables IN MEMORY (tests/ must not import corpus/). Fixture 3 is the
deterministic RED/GREEN: on the OLD id()-keying, distinct late cells survive; on the fix, every
cell is visited and every token is gone. Fixtures 1 and 2 pin the concrete leak surface and the
merged-cell-still-deduped invariant respectively.
"""
from __future__ import annotations

import gc
from io import BytesIO

import docx
from docx import Document

import writer.docx_body as docx_body
from writer.docx_body import _redact_cells, redact_docx_body

# A checksum-VALID legacy SK account (auto=True BANKOVY_UCET); prefix 19 and base 1000002 both
# pass the weighted mod-11 gates in detect.identifiers. Its presence in output = a real leak.
VALID_BANKOVY_UCET = "19-1000002/0900"


def _reopen_from(doc: Document) -> Document:
    """Round-trip a Document through the docx serializer and reopen it (proves the output is a
    structurally valid .docx that docx.Document can parse)."""
    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    return docx.Document(buf)


def _all_cell_text(doc: Document) -> str:
    return "\n".join(
        cell.text for table in doc.tables for row in table.rows for cell in row.cells
    )


def _redact_via_file(doc: Document, tmp_path) -> Document:
    """Run the full public pipeline (redact_docx_body) file->file and reopen the result."""
    src = tmp_path / "in.docx"
    dst = tmp_path / "out.docx"
    doc.save(str(src))
    redact_docx_body(str(src), str(dst))
    return docx.Document(str(dst))


# --------------------------------------------------------------------------- Fixture 1
def test_late_table_cell_bankovy_ucet_not_leaked(tmp_path):
    """A valid-checksum BANKOVY_UCET in a LATE cell (row 5) of a regular 7x2 table must be
    redacted, not skipped. On the old id() key this cell was among those silently skipped."""
    doc = Document()
    table = doc.add_table(rows=7, cols=2)
    for r in range(7):
        for c in range(2):
            # benign filler everywhere except the one late PII cell
            table.cell(r, c).paragraphs[0].add_run(f"cell {r},{c} plain text")
    table.cell(5, 1).paragraphs[0].text = ""
    table.cell(5, 1).paragraphs[0].add_run(VALID_BANKOVY_UCET)

    out = _redact_via_file(doc, tmp_path)

    text = _all_cell_text(out)
    assert VALID_BANKOVY_UCET not in text, f"leaked bank account in output: {text!r}"
    assert "[BANKOVY_UCET]" in text  # positive: it was actually redacted, not just missing


# --------------------------------------------------------------------------- Fixture 2
def test_merged_cell_redacted_exactly_once(tmp_path, monkeypatch):
    """A horizontally-merged (gridSpan) cell yields the SAME <w:tc>/paragraph from every spanned
    access. The dedup must still collapse it to a single visit, and the output must reopen."""
    doc = Document()
    table = doc.add_table(rows=2, cols=2)
    merged = table.cell(0, 0).merge(table.cell(0, 1))
    merged.paragraphs[0].add_run("kontakt@example.com")
    table.cell(1, 0).paragraphs[0].add_run("plain a")
    table.cell(1, 1).paragraphs[0].add_run("plain b")

    # Instrument _redact_paragraph to record every <w:p> ELEMENT it processes. Record the
    # element OBJECT (not id() of it): the transient-proxy id() is exactly what CPython
    # recycles, so an id()-based recorder would itself alias distinct paragraphs. Holding the
    # element objects alive in `visited` also pins their identity for the assertions below.
    visited: list = []
    orig = docx_body._redact_paragraph

    def spy(paragraph, known_entities):
        visited.append(paragraph._p)
        return orig(paragraph, known_entities)

    monkeypatch.setattr(docx_body, "_redact_paragraph", spy)

    _redact_cells(table, None)

    merged_p = merged.paragraphs[0]._p
    assert visited.count(merged_p) == 1, "merged cell processed more than once"
    # every distinct <w:p> visited exactly once (merged cell counts once, not per span)
    assert len(visited) == len(set(visited))

    reopened = _reopen_from(doc)
    assert "kontakt@example.com" not in _all_cell_text(reopened)
    # Check the merged cell itself (row.cells surfaces a span twice, so _all_cell_text would
    # double-count its text); the merged cell holds exactly one label after a single redaction.
    merged_out = reopened.tables[0].cell(0, 0)
    assert merged_out.text == "[EMAIL]", f"expected one label, got: {merged_out.text!r}"


# --------------------------------------------------------------------------- Fixture 3
def test_every_cell_visited_no_token_skipped(tmp_path):
    """Visitation invariant: a regular table whose every cell holds a UNIQUE detectable token.
    After redact, NO token may survive -> no cell was skipped. This is the deterministic
    RED (old id() key leaves late cells' tokens present) / GREEN (fix removes all) case."""
    rows, cols = 7, 2
    doc = Document()
    table = doc.add_table(rows=rows, cols=cols)
    tokens = []
    k = 0
    for r in range(rows):
        for c in range(cols):
            tok = f"tokencell{k}@example.com"
            tokens.append(tok)
            table.cell(r, c).paragraphs[0].add_run(tok)
            k += 1

    # Force the proxy churn / GC that triggers the address-reuse collision on the old key.
    gc.collect()
    _redact_cells(table, None)
    gc.collect()

    survivors = [t for t in tokens if any(t in cell.text for row in table.rows for cell in row.cells)]
    assert survivors == [], f"cells skipped, tokens survived: {survivors}"

    # And the redacted doc still reopens.
    reopened = _reopen_from(doc)
    assert "@example.com" not in _all_cell_text(reopened)

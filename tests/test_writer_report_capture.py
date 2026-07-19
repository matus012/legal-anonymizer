"""Report-capture unit tests: the DOCX writer must, as a pure side-effect of the redaction
pass, record enough to emit a per-document <name>_report.txt in a LATER round. THIS round
captures only — it writes no file and changes NO redaction behaviour or label string.

Two captures are pinned here, both driven off the SAME detect() call the redaction path
already makes on the post-strip paragraph text:

  * REDACTIONS: every kept (auto=True) span -> (label, location, surface), for EVERY
    occurrence including repeats (never deduped), tagged with the caller's location.
  * LOW-CONFIDENCE: every auto=False Candidate detect() returns -> (location, type, surface),
    WITHOUT redacting it and WITHOUT minting a label.

Fixtures are hand-built in memory with python-docx (docx.Document().add_paragraph). NO corpus
import. detect() and python-docx are the only heavy imports. Each test drives _redact_paragraph
directly with an explicit ``location`` and a shared LabelMap so the capture side-effects are
observed without assembling a whole document.

RED-vs-GREEN: at HEAD, LabelMap has no occurrences/low_confidence structures and
_redact_paragraph takes no ``location`` param, so every test errors/fails; only the capture
implementation makes them pass — and it must do so WITHOUT redacting the low-confidence literal
(test 3) or perturbing the first-seen numbering (test 4).
"""
from __future__ import annotations

from docx import Document

from writer.docx_body import _redact_paragraph
from writer.labelmap import LabelMap


def _para(text: str):
    """A single hand-built paragraph (one run); paragraph.text == text before redaction."""
    return Document().add_paragraph(text)


def test_occurrence_recorded_for_each_location() -> None:
    """1. Same known entity in two locations -> occurrences[label] has exactly 2 tuples,
    with the correct per-call location tags and count 2 (one document-global label)."""
    lm = LabelMap(["Novak"])
    p_body = _para("Novak")
    p_header = _para("Novak")

    _redact_paragraph(p_body, ["Novak"], lm, location="body")
    _redact_paragraph(p_header, ["Novak"], lm, location="header")

    assert list(lm.occurrences.keys()) == ["[MENO_1]"]
    assert lm.occurrences["[MENO_1]"] == [("body", "Novak"), ("header", "Novak")]
    assert len(lm.occurrences["[MENO_1]"]) == 2


def test_repeated_occurrence_same_location_not_deduped() -> None:
    """2. The same name twice in ONE paragraph -> 2 tuples for that label (repeats kept)."""
    lm = LabelMap(["Novak"])
    p = _para("Novak a Novak")

    _redact_paragraph(p, ["Novak"], lm, location="body")

    assert lm.occurrences["[MENO_1]"] == [("body", "Novak"), ("body", "Novak")]


def test_low_confidence_captured_without_redacting_or_labelling() -> None:
    """3. A checksum-INVALID rodne cislo (detect -> RODNE_CISLO, auto=False) is captured in
    low_confidence with its location tag and type, is NOT redacted (literal survives in the
    paragraph text), and mints NO label."""
    lm = LabelMap([])
    p = _para("Rodne cislo 900101/0000 tu.")

    _redact_paragraph(p, [], lm, location="footnote")

    assert lm.low_confidence == [("footnote", "RODNE_CISLO", "900101/0000")]
    # left in the output text, unchanged (NOT redacted)
    assert "900101/0000" in p.text
    assert p.text == "Rodne cislo 900101/0000 tu."
    # no label was created for it
    assert lm.occurrences == {}
    assert "[RODNE_CISLO_1]" not in p.text


def test_capture_is_deterministic_across_runs() -> None:
    """4. First-seen label numbering + occurrence order are stable across two runs on
    identical input."""

    def run() -> dict:
        lm = LabelMap(["Novak", "Horak"])
        p1 = _para("Novak a Horak")
        p2 = _para("Novak zas")
        _redact_paragraph(p1, ["Novak", "Horak"], lm, location="body")
        _redact_paragraph(p2, ["Novak", "Horak"], lm, location="header")
        return dict(lm.occurrences)

    first = run()
    second = run()

    assert first == second
    assert first["[MENO_1]"] == [("body", "Novak"), ("header", "Novak")]
    assert first["[MENO_2]"] == [("body", "Horak")]

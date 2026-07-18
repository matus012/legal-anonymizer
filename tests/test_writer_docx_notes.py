"""W3 (context.md §10): redact auto-detected PII in the three note OPC parts Word keeps
OUTSIDE document.xml — footnotes.xml, endnotes.xml and comments.xml.

These fixtures are built entirely in-memory (this module must NOT import corpus/): a minimal
python-docx Document, plus footnotes / endnotes / comments parts hand-attached as generic
``Part`` objects with the correct content types and relationship types. The crucial part-type
asymmetry — footnotes/endnotes reopen as blob-backed generic Parts, comments reopens as an
element-backed CommentsPart — only materialises AFTER a save()+reopen round-trip, and
``redact_docx_body`` takes file paths, so every test writes the fixture to disk first and
redacts path->path exactly as production does.

RED/GREEN: run this file against the pre-W3 writer and the three note assertions FAIL because
each target survives untouched into the output part (the leak message prints the surviving
bytes). After wiring W3 into redact_docx_body they pass: target gone, type label present,
separator entries intact, output reopens clean.
"""
from __future__ import annotations

import zipfile

from docx import Document
from docx.opc.constants import RELATIONSHIP_TYPE as RT
from docx.opc.packuri import PackURI
from docx.opc.part import Part
from docx.oxml.ns import nsdecls

from writer.docx_body import redact_docx_body

# Auto-detected surfaces, checksum-verified against detect() (see probe): RODNE_CISLO and ICO
# carry valid checksums -> auto=True; the IBAN passes mod-97 -> auto=True. If any of these
# ever flipped to review-bucket (auto=False) the corresponding assertion would (correctly)
# stop expecting redaction — but per the pinned spec these three are auto.
RC = "785907/0076"          # -> [RODNE_CISLO]
ICO = "48098191"            # -> [ICO]
IBAN = "SK3411000000002600000017"  # mod-97 valid -> [IBAN]

_FN_CT = "application/vnd.openxmlformats-officedocument.wordprocessingml.footnotes+xml"
_EN_CT = "application/vnd.openxmlformats-officedocument.wordprocessingml.endnotes+xml"
_CM_CT = "application/vnd.openxmlformats-officedocument.wordprocessingml.comments+xml"

# A footnote/endnote separator entry: a <w:p> whose only run holds <w:separator/> and NO
# <w:t>. It must pass through the redactor as a no-op and survive verbatim in the output.
_FN_SEPS = (
    '<w:footnote w:type="separator" w:id="-1"><w:p><w:r><w:separator/></w:r></w:p></w:footnote>'
    '<w:footnote w:type="continuationSeparator" w:id="0"><w:p><w:r>'
    "<w:continuationSeparator/></w:r></w:p></w:footnote>"
)
_EN_SEPS = (
    '<w:endnote w:type="separator" w:id="-1"><w:p><w:r><w:separator/></w:r></w:p></w:endnote>'
    '<w:endnote w:type="continuationSeparator" w:id="0"><w:p><w:r>'
    "<w:continuationSeparator/></w:r></w:p></w:endnote>"
)


def _attach(doc, name: str, ct: str, xml: str, rt) -> None:
    """Attach ``xml`` as a generic Part at /word/<name> and relate it from the main part.
    On reopen python-docx maps the content type to a part class: comments+xml has a
    registered CommentsPart, footnotes/endnotes+xml have none and stay generic blob Parts."""
    part = Part(PackURI(f"/word/{name}"), ct, xml.encode("utf-8"), doc.part.package)
    doc.part.relate_to(part, rt)


def _build_fixture(path) -> str:
    """A minimal docx with all three note parts, saved to ``path`` (str). Returns the path."""
    doc = Document()
    doc.add_paragraph("Telo dokumentu bez PII.")

    fn_body = (
        f'<w:footnote w:id="2"><w:p><w:r>'
        f'<w:t xml:space="preserve">Poznamka pod ciarou, rodne cislo {RC} osoby.</w:t>'
        f"</w:r></w:p></w:footnote>"
    )
    en_body = (
        f'<w:endnote w:id="2"><w:p><w:r>'
        f'<w:t xml:space="preserve">Vysvetlivka: subjekt s ICO {ICO} v registri.</w:t>'
        f"</w:r></w:p></w:endnote>"
    )
    cm_body = (
        f'<w:comment w:id="1" w:author="Advokat" w:date="2024-01-01T00:00:00Z" '
        f'w:initials="AK"><w:p><w:r>'
        f'<w:t xml:space="preserve">Overit platbu na ucet {IBAN} prosim.</w:t>'
        f"</w:r></w:p></w:comment>"
    )

    _attach(doc, "footnotes.xml", _FN_CT, f'<w:footnotes {nsdecls("w")}>{_FN_SEPS}{fn_body}</w:footnotes>', RT.FOOTNOTES)
    _attach(doc, "endnotes.xml", _EN_CT, f'<w:endnotes {nsdecls("w")}>{_EN_SEPS}{en_body}</w:endnotes>', RT.ENDNOTES)
    _attach(doc, "comments.xml", _CM_CT, f'<w:comments {nsdecls("w")}>{cm_body}</w:comments>', RT.COMMENTS)

    doc.save(str(path))
    return str(path)


def _part_xml(docx_path: str, member: str) -> str:
    """Raw, bytes-exact XML string of a zip member (NBSP-sensitive substring checks)."""
    with zipfile.ZipFile(docx_path) as z:
        return z.read(member).decode("utf-8")


def _redacted(tmp_path) -> str:
    in_path = _build_fixture(tmp_path / "in.docx")
    # Pre-change sanity: the raw fixture really does contain every target, so a later
    # "target absent" assertion is a genuine RED->GREEN signal, not a vacuous pass.
    assert RC in _part_xml(in_path, "word/footnotes.xml")
    assert ICO in _part_xml(in_path, "word/endnotes.xml")
    assert IBAN in _part_xml(in_path, "word/comments.xml")

    out_path = str(tmp_path / "out.docx")
    redact_docx_body(in_path, out_path)
    return out_path


def test_footnote_pii_redacted(tmp_path):
    out = _redacted(tmp_path)
    xml = _part_xml(out, "word/footnotes.xml")
    assert RC not in xml, f"LEAKED footnote RC: {RC!r} survived in {xml!r}"
    assert "[RODNE_CISLO]" in xml
    # separator entries untouched
    assert "<w:separator/>" in xml and "<w:continuationSeparator/>" in xml
    Document(out)  # reopens without corruption


def test_endnote_pii_redacted(tmp_path):
    out = _redacted(tmp_path)
    xml = _part_xml(out, "word/endnotes.xml")
    assert ICO not in xml, f"LEAKED endnote ICO: {ICO!r} survived in {xml!r}"
    assert "[ICO]" in xml
    assert "<w:separator/>" in xml and "<w:continuationSeparator/>" in xml
    Document(out)


def test_comment_pii_redacted(tmp_path):
    out = _redacted(tmp_path)
    xml = _part_xml(out, "word/comments.xml")
    assert IBAN not in xml, f"LEAKED comment IBAN: {IBAN!r} survived in {xml!r}"
    assert "[IBAN]" in xml
    Document(out)


def test_comment_element_path_persists_through_save(tmp_path):
    """Comments reopen as a CommentsPart whose .blob RE-SERIALIZES from .element, so W3 must
    mutate part.element in place; a ._blob reassignment would be silently discarded. Reading
    the SAVED output zip (not the in-memory tree) proves the element mutation actually
    persisted through doc.save(), guarding the '_blob-reassign-is-dead' trap."""
    out = _redacted(tmp_path)
    # persisted on disk:
    xml = _part_xml(out, "word/comments.xml")
    assert IBAN not in xml and "[IBAN]" in xml
    # and the reopened object sees the same:
    doc = Document(out)
    (comments_part,) = [
        r.target_part for r in doc.part.rels.values() if r.reltype.endswith("comments")
    ]
    reopened = comments_part.blob.decode("utf-8")
    assert IBAN not in reopened and "[IBAN]" in reopened
    # author attribute is W4 scope — W3 leaves it:
    assert 'w:author="Advokat"' in xml


def test_body_still_clean_and_all_parts_reopen(tmp_path):
    """End-to-end: after W3 the whole document still reopens and the non-note body is intact."""
    out = _redacted(tmp_path)
    doc = Document(out)
    assert doc.paragraphs[0].text == "Telo dokumentu bez PII."
    for member in ("word/footnotes.xml", "word/endnotes.xml", "word/comments.xml"):
        _part_xml(out, member)  # member exists / readable

"""W5b-2 Step 1: the per-document <stem>_report.txt.

Two layers, both pinned here:

  * build_report(occurrences, low_confidence) -> str is PURE (no I/O). It formats the two
    report sections deterministically: REDACTED rows sorted by (TYPE, numeric N) so [MENO_2]
    precedes [MENO_10] and types group, and a LOW CONFIDENCE section in capture (list) order.
    TYPE is recovered from a "[TYPE_N]" label by an rsplit-on-"_"-ONCE, so a multi-underscore
    type like "[RODNE_CISLO_1]" reports TYPE "RODNE_CISLO", not "RODNE".
  * write_report / redact_docx_body wires the file next to out_path (<stem>_report.txt) as an
    ADDITIVE side-effect: it must not change the redacted .docx bytes, and the low-confidence
    literal stays UNREDACTED in the doc while being listed (only) in the low-confidence section.

RED at HEAD: writer.report does not exist, so this module fails to import (collection error)
until build_report / write_report are implemented. The section markers asserted below are the
literal ones build_report emits; the expected row strings are re-derived here from the SAME
" | "-join rule the builder uses, never pasted from program output.
"""
from __future__ import annotations

from docx import Document

from writer.docx_body import redact_docx_body
from writer.report import build_report

# Literal section markers build_report emits; the test splits/locates on these, not on layout.
_REDACTED_MARKER = "[REDACTED]"
_LOWCONF_MARKER = "[LOW CONFIDENCE / NOT REDACTED]"
# Stable liability substring build_report puts in its header (context.md "Liability posture").
_LIABILITY_SUBSTR = "does NOT certify"


def _line_containing(text: str, needle: str) -> str:
    """The single report line that contains ``needle`` (fails loudly if 0 or >1)."""
    hits = [ln for ln in text.splitlines() if needle in ln]
    assert len(hits) == 1, f"expected exactly one line with {needle!r}, got {hits!r}"
    return hits[0]


def test_build_report_content() -> None:
    """A (pure): ordering, count, sorted-set locations, rsplit-once TYPE, low-conf list order,
    and the liability framing — all off hand-built captures, no redaction run, no disk."""
    occurrences = {
        "[MENO_1]": [("body", "Novak"), ("header", "Novak"), ("footnote", "Novaka")],
        "[MENO_2]": [("body", "Horak")],
        "[RODNE_CISLO_1]": [("body", "880101/0000")],
    }
    low_confidence = [
        ("footnote", "RODNE_CISLO", "900101/0000"),
        ("body", "DIC", "2502182454"),
    ]

    report = build_report(occurrences, low_confidence)

    # Liability framing sentence present (header block).
    assert _LIABILITY_SUBSTR in report

    # Row strings re-derived from the SAME rule the builder uses (TYPE | label | count | locations).
    row_meno1 = " | ".join(["MENO", "[MENO_1]", "3", "body, footnote, header"])
    row_meno2 = " | ".join(["MENO", "[MENO_2]", "1", "body"])
    row_rc1 = " | ".join(["RODNE_CISLO", "[RODNE_CISLO_1]", "1", "body"])
    assert row_meno1 in report  # count 3 (repeats), locations = sorted set
    assert row_meno2 in report
    assert row_rc1 in report

    # Ordering: (TYPE, numeric N) -> MENO_1 < MENO_2 < RODNE_CISLO_1.
    i_meno1 = report.index("[MENO_1]")
    i_meno2 = report.index("[MENO_2]")
    i_rc1 = report.index("[RODNE_CISLO_1]")
    assert i_meno1 < i_meno2 < i_rc1

    # TYPE parsed via rsplit-once: [RODNE_CISLO_1]'s row starts with the full type, not "RODNE".
    rc_line = _line_containing(report, "[RODNE_CISLO_1]")
    assert rc_line.startswith("RODNE_CISLO | ")

    # Low-confidence section, in capture (list) order: RODNE_CISLO before DIC.
    lowconf_part = report.split(_LOWCONF_MARKER, 1)[1]
    lc_rc = " | ".join(["RODNE_CISLO", "900101/0000", "footnote"])
    lc_dic = " | ".join(["DIC", "2502182454", "body"])
    assert lc_rc in lowconf_part
    assert lc_dic in lowconf_part
    assert lowconf_part.index(lc_rc) < lowconf_part.index(lc_dic)

    # Empty low_confidence prints an explicit "(none)" line, not a blank section.
    empty_report = build_report({"[MENO_1]": [("body", "Novak")]}, [])
    assert "(none)" in empty_report.split(_LOWCONF_MARKER, 1)[1]


def test_report_written_next_to_out(tmp_path) -> None:
    """B (integration): redact_docx_body writes <stem>_report.txt next to out_path as an additive
    side-effect. The known-entity MENO is redacted+listed; the checksum-invalid RC (auto=False)
    is NOT redacted, absent from the REDACTED section, present in the low-confidence section, and
    still literally present in the redacted .docx text."""
    in_path = tmp_path / "in.docx"
    out_path = tmp_path / "out_anon.docx"
    doc = Document()
    doc.add_paragraph("Novak Rodne cislo 900101/0000 koniec.")
    doc.save(str(in_path))

    redact_docx_body(str(in_path), str(out_path), known_entities=["Novak"])

    report_path = tmp_path / "out_anon_report.txt"
    assert report_path.exists(), "report file must be written next to out_path"
    report = report_path.read_text(encoding="utf-8")

    # Report is not the .docx (additive, no collision with out_anon.docx).
    assert out_path.exists()
    assert report_path != out_path

    # MENO row present; low-confidence RODNE_CISLO line present.
    assert "[MENO_1]" in report
    assert _line_containing(report, "[MENO_1]").startswith("MENO | ")
    lowconf_part = report.split(_LOWCONF_MARKER, 1)[1]
    assert "RODNE_CISLO" in lowconf_part
    assert "900101/0000" in lowconf_part

    # The checksum-invalid literal is ABSENT from the REDACTED section but PRESENT in low-conf.
    redacted_part = report.split(_LOWCONF_MARKER, 1)[0]
    assert "900101/0000" not in redacted_part
    assert "900101/0000" in lowconf_part

    # The redacted .docx reopens; MENO label is in its text; the RC literal survives unredacted.
    reopened = Document(str(out_path))
    body_text = "\n".join(p.text for p in reopened.paragraphs)
    assert "[MENO_1]" in body_text
    assert "900101/0000" in body_text

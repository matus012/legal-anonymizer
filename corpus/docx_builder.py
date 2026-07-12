"""Low-level DOCX assembly with every seeded DOCX failure mode (context.md §7, §10).

Beyond what python-docx offers, this hand-builds the parts a naive ``run.text`` scan misses:
split runs, footnotes/endnotes/comments (separate OPC parts), VML textboxes, tracked changes
(``w:ins`` / ``w:del`` — deleted text persists), and metadata in core.xml / app.xml.

Every placed :class:`PiiSpec` is recorded in the shared :class:`Recorder` with the exact
``surface_part`` so the ground truth never lies about where a string lives.
"""
from __future__ import annotations

import datetime as _dt
import os
import random
import zipfile
from pathlib import Path
from xml.sax.saxutils import escape

import docx
from docx.opc.constants import RELATIONSHIP_TYPE as RT
from docx.opc.packuri import PackURI
from docx.opc.part import Part
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls, qn

from .groundtruth import PiiSpec, Recorder

W = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"

Item = "str | PiiSpec"

_SPLIT_PROB = 0.30  # fraction of PII surfaces split across runs (context.md §7 failure mode)

# Fixed timestamps so the same seed yields byte-identical files (no wall-clock in the output).
_FIXED_DT = _dt.datetime(2025, 1, 1, 0, 0, 0)
_FIXED_ZIP_TIME = (2025, 1, 1, 0, 0, 0)


def _pin_zip_timestamps(path: Path) -> None:
    """Rewrite every zip entry with a fixed mtime — otherwise DOCX save stamps wall-clock time
    into each entry and the same seed produces different bytes."""
    with zipfile.ZipFile(str(path)) as zin:
        infos = zin.infolist()
        blobs = {i.filename: zin.read(i.filename) for i in infos}
    tmp = str(path) + ".tmp"
    with zipfile.ZipFile(tmp, "w", zipfile.ZIP_DEFLATED) as zout:
        for i in infos:
            zi = zipfile.ZipInfo(i.filename, date_time=_FIXED_ZIP_TIME)
            zi.compress_type = i.compress_type
            zi.external_attr = i.external_attr
            zi.internal_attr = i.internal_attr
            zi.create_system = i.create_system
            zout.writestr(zi, blobs[i.filename])
    os.replace(tmp, str(path))


def _text_of(items: list) -> str:
    return "".join(it if isinstance(it, str) else it.surface for it in items)


def _run_xml(text: str) -> str:
    return f'<w:r><w:t xml:space="preserve">{escape(text)}</w:t></w:r>'


class DocxBuilder:
    def __init__(self, recorder: Recorder, rng=None):
        self.rec = recorder
        self.rng = rng if rng is not None else random.Random(0)
        self.doc = docx.Document()
        self._para_idx = -1
        self._fn_id = 0
        self._en_id = 0
        self._cmt_id = -1
        self._footnotes: list[str] = []
        self._endnotes: list[str] = []
        self._comments: list[str] = []

    # ---------------------------------------------------------------- body helpers
    def heading(self, text: str, level: int = 1) -> None:
        self.doc.add_heading(text, level=level)

    def _record_items(self, items: list, part: str, detail: dict) -> None:
        for it in items:
            if isinstance(it, PiiSpec):
                self.rec.record(it, part=part, detail=detail)

    def _split_surface(self, s: str) -> list[str]:
        """~30% of the time, cut ``s`` into 2-3 pieces at random offsets (the split-run trap)."""
        if len(s) < 4 or self.rng.random() >= _SPLIT_PROB:
            return [s]
        n = min(self.rng.choice((2, 3)), len(s))
        cuts = sorted(self.rng.sample(range(1, len(s)), n - 1))
        parts, prev = [], 0
        for c in cuts:
            parts.append(s[prev:c])
            prev = c
        parts.append(s[prev:])
        return parts

    def _render_and_record(self, paragraph, items: list, part: str, detail: dict) -> None:
        """Add runs for ``items``; each PiiSpec surface is sometimes split across several runs
        so a regex over individual ``run.text`` finds nothing (context.md §7, §10). Ground
        truth still records the full logical surface — that is the whole point of the trap."""
        for it in items:
            if isinstance(it, str):
                if it:
                    paragraph.add_run(it)
            else:
                chunks = self._split_surface(it.surface)
                for ch in chunks:
                    paragraph.add_run(ch)
                d = dict(detail)
                if len(chunks) > 1:
                    d["split_runs"] = True
                self.rec.record(it, part=part, detail=d)

    def paragraph(self, items: list, *, style: str | None = None) -> None:
        """A body paragraph built from plain strings and PiiSpecs (PII sometimes split-run)."""
        self._para_idx += 1
        p = self.doc.add_paragraph(style=style)
        self._render_and_record(p, items, "body", {"paragraph_index": self._para_idx})

    def split_run_paragraph(self, prefix: str, spec: PiiSpec, suffix: str, chunks: int = 3) -> None:
        """Place ``spec.surface`` split across several runs — the classic silent-miss case.

        A regex over individual ``run.text`` values finds nothing; only reconstructed
        paragraph text reveals the name (context.md §7, §10).
        """
        self._para_idx += 1
        p = self.doc.add_paragraph()
        if prefix:
            p.add_run(prefix)
        s = spec.surface
        n = max(2, min(chunks, len(s)))
        step = max(1, len(s) // n)
        for i in range(0, len(s), step):
            p.add_run(s[i : i + step])
        if suffix:
            p.add_run(suffix)
        self.rec.record(
            spec, part="body", detail={"paragraph_index": self._para_idx, "split_runs": True}
        )

    def table(self, rows: list[list[list]]) -> None:
        """``rows[r][c]`` is an items list for that cell."""
        n_cols = max(len(r) for r in rows)
        t = self.doc.add_table(rows=len(rows), cols=n_cols)
        t.style = "Table Grid"
        for r, row in enumerate(rows):
            for c, cell_items in enumerate(row):
                cell = t.cell(r, c)
                self._render_and_record(cell.paragraphs[0], cell_items, "table_cell",
                                        {"row": r, "col": c})

    # ---------------------------------------------------------------- header / footer
    def header(self, items: list) -> None:
        para = self.doc.sections[0].header.paragraphs[0]
        para.text = _text_of(items)
        self._record_items(items, "header", {})

    def footer(self, items: list) -> None:
        para = self.doc.sections[0].footer.paragraphs[0]
        para.text = _text_of(items)
        self._record_items(items, "footer", {})

    # ---------------------------------------------------------------- footnotes / endnotes
    def footnote(self, anchor_para: str, items: list) -> None:
        self._fn_id += 1
        fid = self._fn_id
        self._para_idx += 1
        p = self.doc.add_paragraph(anchor_para)
        run = p.add_run()
        ref = parse_xml(f'<w:footnoteReference {nsdecls("w")} w:id="{fid}"/>')
        run._r.append(ref)
        self._footnotes.append(
            f'<w:footnote w:id="{fid}"><w:p>{_run_xml(_text_of(items))}</w:p></w:footnote>'
        )
        self._record_items(items, "footnote", {"footnote_id": fid})

    def endnote(self, anchor_para: str, items: list) -> None:
        self._en_id += 1
        eid = self._en_id
        self._para_idx += 1
        p = self.doc.add_paragraph(anchor_para)
        run = p.add_run()
        ref = parse_xml(f'<w:endnoteReference {nsdecls("w")} w:id="{eid}"/>')
        run._r.append(ref)
        self._endnotes.append(
            f'<w:endnote w:id="{eid}"><w:p>{_run_xml(_text_of(items))}</w:p></w:endnote>'
        )
        self._record_items(items, "endnote", {"endnote_id": eid})

    # ---------------------------------------------------------------- comments
    def comment(self, anchor_text: str, items: list, author: str = "Advokát") -> None:
        self._cmt_id += 1
        cid = self._cmt_id
        self._para_idx += 1
        p = self.doc.add_paragraph()
        start = parse_xml(f'<w:commentRangeStart {nsdecls("w")} w:id="{cid}"/>')
        end = parse_xml(f'<w:commentRangeEnd {nsdecls("w")} w:id="{cid}"/>')
        p._p.append(start)
        p.add_run(anchor_text)
        p._p.append(end)
        ref_run = parse_xml(
            f'<w:r {nsdecls("w")}><w:commentReference w:id="{cid}"/></w:r>'
        )
        p._p.append(ref_run)
        self._comments.append(
            f'<w:comment w:id="{cid}" w:author="{escape(author)}" w:date="2024-01-01T00:00:00Z" '
            f'w:initials="AK"><w:p>{_run_xml(_text_of(items))}</w:p></w:comment>'
        )
        self._record_items(items, "comment", {"comment_id": cid})

    # ---------------------------------------------------------------- textbox (VML)
    def textbox(self, items: list) -> None:
        self._para_idx += 1
        p = self.doc.add_paragraph()
        ns = f'xmlns:w="{W}" xmlns:v="urn:schemas-microsoft-com:vml"'
        xml = (
            f'<w:r {ns}><w:pict><v:shape type="#_x0000_t202" '
            f'style="width:240pt;height:48pt"><v:textbox><w:txbxContent>'
            f'<w:p>{_run_xml(_text_of(items))}</w:p>'
            f"</w:txbxContent></v:textbox></v:shape></w:pict></w:r>"
        )
        p._p.append(parse_xml(xml))
        self._record_items(items, "textbox", {})

    # ---------------------------------------------------------------- tracked changes
    def tracked_change(
        self, lead: str, ins_items: list, del_items: list, author: str = "Advokát"
    ) -> None:
        """One paragraph carrying a ``w:ins`` and a ``w:del`` — deleted text stays in the file."""
        self._para_idx += 1
        p = self.doc.add_paragraph()
        if lead:
            p.add_run(lead)
        meta = 'w:author="%s" w:date="2024-01-01T00:00:00Z"' % escape(author)
        ins_runs = "".join(_run_xml(_text_of([it])) for it in ins_items)
        p._p.append(parse_xml(f'<w:ins {nsdecls("w")} w:id="9001" {meta}>{ins_runs}</w:ins>'))
        del_runs = "".join(
            f'<w:r><w:delText xml:space="preserve">{escape(_text_of([it]))}</w:delText></w:r>'
            for it in del_items
        )
        p._p.append(parse_xml(f'<w:del {nsdecls("w")} w:id="9002" {meta}>{del_runs}</w:del>'))
        self._record_items(ins_items, "tracked_change_ins", {"paragraph_index": self._para_idx})
        self._record_items(del_items, "tracked_change_del", {"paragraph_index": self._para_idx})

    # ---------------------------------------------------------------- metadata
    def set_metadata(self, author_spec: PiiSpec, company_spec: PiiSpec | None = None) -> None:
        cp = self.doc.core_properties
        cp.author = author_spec.surface
        cp.last_modified_by = author_spec.surface
        self.rec.record(author_spec, part="metadata_core", detail={"field": "author"})
        if company_spec is not None:
            self._set_company(company_spec.surface)
            self.rec.record(company_spec, part="metadata_app", detail={"field": "Company"})

    def _set_company(self, company: str) -> None:
        for part in self.doc.part.package.iter_parts():
            if part.partname == "/docProps/app.xml":
                blob = part._blob.decode("utf-8")
                tag = f"<Company>{escape(company)}</Company>"
                if "<Company>" in blob:
                    import re

                    blob = re.sub(r"<Company>.*?</Company>", tag, blob)
                elif "<Company/>" in blob:
                    blob = blob.replace("<Company/>", tag)
                else:
                    blob = blob.replace("</Properties>", tag + "</Properties>")
                part._blob = blob.encode("utf-8")
                return

    # ---------------------------------------------------------------- save
    def _attach_notes_part(self, tag: str, ct_leaf: str, rt, seps: str, bodies: list[str], fname: str):
        if not bodies:
            return
        xml = f'<w:{tag} {nsdecls("w")}>{seps}{"".join(bodies)}</w:{tag}>'
        ct = f"application/vnd.openxmlformats-officedocument.wordprocessingml.{ct_leaf}+xml"
        part = Part(PackURI(f"/word/{fname}"), ct, xml.encode("utf-8"), self.doc.part.package)
        self.doc.part.relate_to(part, rt)

    def save(self, path: Path) -> None:
        fn_seps = (
            '<w:footnote w:type="separator" w:id="-1"><w:p><w:r><w:separator/></w:r></w:p></w:footnote>'
            '<w:footnote w:type="continuationSeparator" w:id="0"><w:p><w:r><w:continuationSeparator/>'
            "</w:r></w:p></w:footnote>"
        )
        en_seps = (
            '<w:endnote w:type="separator" w:id="-1"><w:p><w:r><w:separator/></w:r></w:p></w:endnote>'
            '<w:endnote w:type="continuationSeparator" w:id="0"><w:p><w:r><w:continuationSeparator/>'
            "</w:r></w:p></w:endnote>"
        )
        self._attach_notes_part("footnotes", "footnotes", RT.FOOTNOTES, fn_seps, self._footnotes, "footnotes.xml")
        self._attach_notes_part("endnotes", "endnotes", RT.ENDNOTES, en_seps, self._endnotes, "endnotes.xml")
        if self._comments:
            xml = f'<w:comments {nsdecls("w")}>{"".join(self._comments)}</w:comments>'
            ct = "application/vnd.openxmlformats-officedocument.wordprocessingml.comments+xml"
            part = Part(PackURI("/word/comments.xml"), ct, xml.encode("utf-8"), self.doc.part.package)
            self.doc.part.relate_to(part, RT.COMMENTS)
        cp = self.doc.core_properties  # pin document dates
        cp.created = _FIXED_DT
        cp.modified = _FIXED_DT
        self.doc.save(str(path))
        _pin_zip_timestamps(path)

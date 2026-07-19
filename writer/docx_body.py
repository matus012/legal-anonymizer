"""W1 (context.md §10): redact auto-detected PII in a DOCX's top-level body paragraphs.

The hard part is not *finding* PII — detect() does that on reconstructed paragraph text —
it is putting the redaction back onto the <w:r> run sequence. Word splits a single logical
surface across arbitrary runs, sometimes MID-TOKEN ('Luc'|'ie Molnáro'|'vej') and sometimes
MID-RUN on both ends ('Zmluva medzi Ján'|'om Novákom dnes'), and each run may carry its own
<w:rPr> formatting. So a detect() span [start,end) over the reconstructed text has to be
mapped back to exact character offsets inside runs, the covered characters removed, a single
type label inserted at the span start, and any boundary run split into fragments that each
keep a clone of the original run's <w:rPr>.

W2 (context.md §10) extends coverage past doc.paragraphs to every OTHER place Word hides a
<w:p>: table cells, header/footer paragraphs, header/footer tables, and VML textboxes (in
the body and in header/footer parts). All of them are ordinary <w:p> once located, so the
SAME run-remap core (_redact_paragraph) is reused verbatim — nothing about the run-splitting
is re-implemented.

W3 (context.md §10) closes the last <w:p> locations, the three note parts Word keeps OUTSIDE
document.xml as separate OPC parts: footnotes.xml, endnotes.xml and comments.xml. They carry a
part-type asymmetry (see _redact_notes_part) but every note is an ordinary <w:p> once located,
so the SAME core is reused again. W3 does NOT scrub the comment w:author attribute — that is
W4 metadata scope. Labels are type-only ("[MENO]"); per-entity numbering ("[MENO_1]") is W5.
The input file is never modified — output is a new file.

W4b (context.md §10) scrubs document METADATA that carries PII: docProps/core.xml properties
(dc:creator, cp:lastModifiedBy, ...), docProps/app.xml's Company/Manager free-text fields, and
the comment w:author/w:initials attributes deferred from W3 (see _redact_notes_part). These are
BLANKED UNCONDITIONALLY BY POSITION — detect() never runs over metadata, since an author or
manager name is PII regardless of whether it matches a detector pattern, and the original value
is never preserved (see _scrub_metadata).
"""
from __future__ import annotations

import re
from copy import deepcopy

from docx import Document
from docx.oxml import parse_xml
from docx.oxml.ns import qn
from docx.text.paragraph import Paragraph
from lxml import etree

from detect.core import detect
from writer.labelmap import LabelMap

# xml:space lives in the reserved XML namespace, which is not in python-docx's nsmap, so it
# cannot go through qn(); set it by its literal Clark-notation name.
_XML_SPACE = "{http://www.w3.org/XML/1998/namespace}space"


def _strip_tracked_changes(root) -> None:
    """W4a (context.md §10): accept ALL Word tracked revisions inside ``root`` (a w:document /
    header / footer / notes tree) BEFORE any redaction runs. The structure is ASYMMETRIC
    (verified against corpus bytes, kupna_zmluva_000.docx):

    * <w:ins ...><w:r><w:t>TEXT</w:t></w:r></w:ins> — accepted INSERTED content: promote the
      ins's children into its parent at the ins's position, then drop the now-empty ins. The
      runs SURVIVE as ordinary body runs and remain subject to the normal redaction passes.
    * <w:del ...><w:r><w:delText>TEXT</w:delText></w:r></w:del> — DELETED content that is still
      PHYSICALLY in the file: remove the whole subtree. Getting these backwards either drops
      accepted text or leaks deleted PII. w:delText only ever lives inside w:del, so once every
      w:del is gone no w:delText remains.

    Every w:del is stripped first (removing its whole subtree), then w:ins is re-found on the
    mutated tree — so content that was inserted-then-deleted (a w:del nested in a w:ins, or an
    ins nested in a del) resolves correctly with no dangling detached-parent inserts. Both lists
    are materialised (findall + list) so the tree is never mutated under a live iterator."""
    for dele in list(root.findall(".//" + qn("w:del"))):
        parent = dele.getparent()
        if parent is not None:
            parent.remove(dele)

    for ins in list(root.findall(".//" + qn("w:ins"))):
        parent = ins.getparent()
        if parent is None:
            continue
        idx = parent.index(ins)
        # Materialise the children: each insert MOVES the child out of ins into parent, so the
        # live child list shrinks as we go; enumerate over the snapshot keeps order + offsets.
        for offset, child in enumerate(list(ins)):
            parent.insert(idx + offset, child)
        parent.remove(ins)


def _strip_notes_tracked_changes(part) -> None:
    """Accept all tracked revisions in a footnotes/endnotes/comments OPC part, honouring the SAME
    element-vs-blob asymmetry as _redact_notes_part: a CommentsPart exposes a live ``.element``
    that re-serialises into ``.blob`` (mutate it in place); generic footnote/endnote Parts are
    blob-backed with no ``.element`` (parse, strip, reassign ``._blob``)."""
    if hasattr(part, "element") and part.element is not None:
        _strip_tracked_changes(part.element)  # live tree; mutate in place
        return
    tree = parse_xml(part._blob)
    _strip_tracked_changes(tree)
    part._blob = etree.tostring(tree, encoding="UTF-8", standalone=True)


def _rebuild_run(run, fragments: list[tuple[str, str]]) -> None:
    """Replace a single <w:r> with one new <w:r> per fragment, cloning the original run's
    formatting onto each. ``fragments`` is an ordered list of ('t', text) surviving-text and
    ('l', label) replacement pieces. An empty list means the whole run was covered -> remove.
    """
    r_elem = run._r
    parent = r_elem.getparent()
    idx = parent.index(r_elem)

    for offset, (_kind, value) in enumerate(fragments):
        clone = deepcopy(r_elem)  # carries a copy of <w:rPr> so formatting is preserved
        t = clone.find(qn("w:t"))
        if t is None:
            t = clone.makeelement(qn("w:t"), {})
            clone.append(t)
        t.text = value
        # Preserve any leading/trailing whitespace in the fragment (labels have none, but a
        # surviving boundary fragment like ' súhlasí' does).
        t.set(_XML_SPACE, "preserve")
        parent.insert(idx + offset, clone)

    parent.remove(r_elem)


def _redact_paragraph(paragraph, known_entities, labelmap, location: str = "body") -> None:
    """Redact auto PII in one paragraph and record report-capture side-effects tagged with
    ``location`` (one of the fixed vocabulary strings, matching GT surface_part). The default is
    ``"body"``, the safe fallback for an un-tagged call — every real caller now passes
    ``location`` explicitly, including ``_redact_cells`` for body table cells."""
    runs = list(paragraph.runs)
    if not runs:
        return

    # Reconstruct paragraph text and remember which run owns each character offset.
    run_start: list[int] = []
    owner: list[int] = []
    pieces: list[str] = []
    pos = 0
    for ri, r in enumerate(runs):
        run_start.append(pos)
        text = r.text
        pieces.append(text)
        owner.extend([ri] * len(text))
        pos += len(text)
    recon = "".join(pieces)

    # ONE detect() over the post-strip reconstructed text; both captures below read this same
    # result the redaction path uses — never a second detect() over a different tree state.
    detected = detect(recon, known_entities)

    # Low-confidence capture: every auto=False span is recorded for review and left ALONE
    # (not covered, no label). This must happen even when no auto span exists (early-return
    # below), so it precedes the `if not cands` guard.
    for c in detected:
        if not c.auto:
            labelmap.record_low_confidence(location, c.type, c.surface)

    cands = [c for c in detected if c.auto]
    if not cands:
        return

    covered = [False] * len(recon)
    label_at: dict[int, str] = {}
    for c in cands:
        for i in range(c.start, c.end):
            covered[i] = True
        # detect() guarantees non-overlapping spans, so one label per start is unambiguous.
        # The explicit "not in" guard (NOT setdefault) is load-bearing: setdefault would
        # eagerly evaluate labelmap.label_for(c) even when c.start is already present, and
        # label_for mints/caches a number on first sighting — so an already-present start
        # would spuriously bump the per-type counter. Guarding keeps numbering exact.
        if c.start not in label_at:
            label_at[c.start] = labelmap.label_for(c)
        # Redaction capture: record EVERY kept occurrence (all spans, incl. repeats of an
        # already-numbered label). Non-overlapping distinct starts mean label_at[c.start] is
        # this candidate's own label; recording here does not touch counters or the cache.
        labelmap.record_occurrence(label_at[c.start], location, c.surface)

    # Build, per touched run, the ordered surviving-text / label fragments, then rewrite it.
    for ri, r in enumerate(runs):
        rs = run_start[ri]
        re = rs + len(r.text)
        touched = any(covered[i] or i in label_at for i in range(rs, re))
        if not touched:
            continue  # leave the run — and its exact XML — completely alone

        fragments: list[tuple[str, str]] = []
        buf = ""
        for i in range(rs, re):
            if i in label_at:
                if buf:
                    fragments.append(("t", buf))
                    buf = ""
                fragments.append(("l", label_at[i]))
            if not covered[i]:
                buf += recon[i]
        if buf:
            fragments.append(("t", buf))

        _rebuild_run(r, fragments)


def _redact_cells(table, known_entities, labelmap, location: str = "table_cell") -> None:
    """Redact every cell paragraph in ``table``. A merged cell makes several (row, col)
    positions return the SAME <w:tc> element; dedup by tc identity so a shared paragraph is
    processed exactly once (reprocessing is otherwise wasted work and re-runs detect on
    already-labelled text).

    Key the set on the <w:tc> ELEMENT (``cell._tc``) itself, NOT ``id(cell._tc)``. row.cells
    mints a fresh _Cell proxy per row, and each row's proxies are GC'd before the next row's
    are allocated; CPython then reuses the freed address, so a later distinct cell's proxy can
    collide with an earlier cell's id() and be silently skipped (leaking its PII). The lxml
    element is stable for the table's lifetime and hashes/compares by identity, and a genuine
    horizontal span yields the SAME element object from every spanned cell access (python-docx
    _Row.cells builds one _Cell per <w:tc> and yields it grid_span times) — so merged cells
    still collapse to one, while distinct cells never alias. Holding the elements in the set
    also keeps them alive, so no address recycling can occur mid-table."""
    seen: set = set()
    for row in table.rows:
        for cell in row.cells:
            if cell._tc in seen:
                continue
            seen.add(cell._tc)
            for paragraph in cell.paragraphs:
                _redact_paragraph(paragraph, known_entities, labelmap, location)


def _redact_textboxes(element, parent, known_entities, labelmap, location: str = "textbox") -> None:
    """Redact every <w:p> nested in any <w:txbxContent> under ``element`` (a document body
    or a header/footer part element). VML textboxes are unreachable through python-docx's
    paragraph APIs, so each inner <w:p> is wrapped as a Paragraph — whose .runs then expose
    the inner runs — and pushed through the same remap. ``parent`` is only a proxy handle;
    the remap operates on the run elements directly, so its exact value is not load-bearing."""
    for txbx in element.findall(".//" + qn("w:txbxContent")):
        for p_elem in txbx.findall(qn("w:p")):
            _redact_paragraph(Paragraph(p_elem, parent), known_entities, labelmap, location)


def _redact_notes_part(part, known_entities, labelmap, location: str = "footnote") -> None:
    """Redact every <w:p> in a footnotes/endnotes/comments OPC part, honouring the part-type
    asymmetry python-docx exposes on reopen (verified by probe):

    * comments.xml maps to a registered CommentsPart (an XmlPart): its ``.element`` is a live
      tree and its ``.blob`` RE-SERIALIZES from that tree, so a ``._blob`` reassignment would be
      silently discarded. The tree is mutated IN PLACE and doc.save() picks it up.
    * footnotes.xml / endnotes.xml have no registered class and come back as generic blob-backed
      Parts with no ``.element``. Their bytes are parsed, the parsed tree is mutated, and the
      serialized result is written back to ``._blob`` (generic ``Part.blob`` returns ``_blob``,
      so doc.save() persists it).

    Branch on element-vs-blob, never on part name or note id. Each <w:p> — including the
    separator / continuationSeparator entries that carry no <w:t> — is wrapped as a Paragraph so
    ``.runs`` exposes the inner runs, then pushed through the same _redact_paragraph core; an
    empty separator paragraph no-ops (detect() over "" yields nothing)."""
    if hasattr(part, "element") and part.element is not None:
        tree = part.element  # live tree; mutate in place
        for p_elem in tree.findall(".//" + qn("w:p")):
            _redact_paragraph(Paragraph(p_elem, part), known_entities, labelmap, location)
        return

    tree = parse_xml(part._blob)
    for p_elem in tree.findall(".//" + qn("w:p")):
        _redact_paragraph(Paragraph(p_elem, part), known_entities, labelmap, location)
    # Mirror python-docx's own part serialization (UTF-8, standalone declaration).
    part._blob = etree.tostring(tree, encoding="UTF-8", standalone=True)


_APP_XML_PII_TAGS = ("Company", "Manager")


def _scrub_metadata(doc) -> None:
    """W4b (context.md §10): blank the three PII-bearing metadata locations, by position —
    never via detect(), never preserving the original value.

    * docProps/core.xml: exposed through python-docx's ``doc.core_properties`` (each is a
      settable str property backed by a ZeroOrOne element); the API write persists through
      doc.save(). created/modified/revision are dates/ints, not PII, and are left alone.
    * docProps/app.xml: python-docx has no API for this part — it comes back as a generic
      blob-backed Part. Company/Manager are located with a targeted regex on the decoded
      blob so every unrelated tag (HeadingPairs, TitlesOfParts, vt: vectors, the <?xml?>
      declaration) is byte-preserved; only a NON-EMPTY <Tag>...</Tag> is rewritten, so an
      already-empty <Tag/>/<Tag></Tag> is left as-is rather than needlessly touched.
    * word/comments.xml <w:comment w:author=...>: deferred from W3's _redact_notes_part.
      Same element-vs-blob asymmetry applies — a CommentsPart's .blob RE-SERIALIZES from its
      live .element, so the attribute is mutated ON THE ELEMENT, never via a ._blob reassign
      (which would be silently discarded on save). w:id/w:date are not PII and are untouched.
    """
    cp = doc.core_properties
    for attr in ("author", "last_modified_by", "title", "subject", "keywords", "category", "comments"):
        setattr(cp, attr, "")

    for part in doc.part.package.iter_parts():
        if str(part.partname) != "/docProps/app.xml":
            continue
        xml = part._blob.decode("utf-8")
        for tag in _APP_XML_PII_TAGS:
            xml = re.sub(rf"<{tag}>.+?</{tag}>", f"<{tag}></{tag}>", xml, flags=re.DOTALL)
        part._blob = xml.encode("utf-8")
        break

    for rel in doc.part.rels.values():
        if not rel.reltype.endswith("comments"):
            continue
        part = rel.target_part
        if not hasattr(part, "element") or part.element is None:
            continue
        for comment in part.element.findall(".//" + qn("w:comment")):
            if comment.get(qn("w:author")) is not None:
                comment.set(qn("w:author"), "")
            if comment.get(qn("w:initials")) is not None:
                comment.set(qn("w:initials"), "")


def redact_docx_body(
    in_path: str, out_path: str, known_entities: list[str] | None = None
) -> None:
    """Open ``in_path``, redact auto-detected PII across every <w:p> location W2 covers —
    body paragraphs, table cells, header/footer paragraphs, header/footer tables and VML
    textboxes (body + header/footer parts) — and save the result to a NEW file ``out_path``.
    ``in_path`` is never modified."""
    doc = Document(in_path)

    # ONE LabelMap for the whole document, built BEFORE the passes: it is threaded through every
    # _redact_paragraph call so a given entity gets the same [TYPE_N] number everywhere (body,
    # tables, headers/footers, textboxes, notes). First-seen group order == the fixed traversal
    # order of the passes below, so the numbering is deterministic (W5a, context.md §10).
    labelmap = LabelMap(known_entities)

    # 0) W4a: accept ALL tracked revisions FIRST, before any _redact_paragraph call. Deleted
    #    text is then physically gone (it must never reach detect() as a redaction candidate);
    #    inserted text, once unwrapped, is ordinary body text that the passes below still redact
    #    (so an insertion carrying PII is caught by the existing W1-W3 code — no special-casing).
    #    Covers the document element, every section header/footer element, and each note part
    #    (same discovery + blob-vs-element asymmetry as the W3 pass below).
    _strip_tracked_changes(doc.element)
    for section in doc.sections:
        for hf in (section.header, section.footer):
            _strip_tracked_changes(hf._element)
    for rel in doc.part.rels.values():
        rt = rel.reltype
        if rt.endswith("footnotes") or rt.endswith("endnotes") or rt.endswith("comments"):
            _strip_notes_tracked_changes(rel.target_part)

    # 1) top-level body paragraphs (W1). Pass "body" explicitly — _redact_paragraph's default is
    #    now "body" too (see its docstring), but every caller names its location regardless.
    for paragraph in doc.paragraphs:
        _redact_paragraph(paragraph, known_entities, labelmap, "body")

    # 2) body tables' cells.
    for table in doc.tables:
        _redact_cells(table, known_entities, labelmap)

    # 3) header/footer paragraphs + their tables, across every section. The location tag
    #    follows the part (header vs footer), including for tables inside a header/footer.
    for section in doc.sections:
        for hf, loc in ((section.header, "header"), (section.footer, "footer")):
            for paragraph in hf.paragraphs:
                _redact_paragraph(paragraph, known_entities, labelmap, loc)
            for table in hf.tables:
                _redact_cells(table, known_entities, labelmap, loc)

    # 4) VML textboxes anywhere: body part, then every header/footer part.
    _redact_textboxes(doc.element.body, doc, known_entities, labelmap)
    for section in doc.sections:
        for hf in (section.header, section.footer):
            _redact_textboxes(hf._element, hf, known_entities, labelmap)

    # 5) footnotes / endnotes / comments — each a SEPARATE OPC part, not in document.xml (W3).
    #    The location tag follows the note part type.
    _NOTE_LOCATIONS = (("footnotes", "footnote"), ("endnotes", "endnote"), ("comments", "comment"))
    for rel in doc.part.rels.values():
        rt = rel.reltype
        loc = next((location for suffix, location in _NOTE_LOCATIONS if rt.endswith(suffix)), None)
        if loc is not None:
            _redact_notes_part(rel.target_part, known_entities, labelmap, loc)

    # 6) W4b: blank PII-bearing metadata (core.xml properties, app.xml Company/Manager, and
    #    the comment w:author/w:initials deferred from W3) LAST, unconditionally by position.
    _scrub_metadata(doc)

    doc.save(out_path)
